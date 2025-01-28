from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
import calendar
import locale
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup
import os
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
from docx.shared import Inches

#### PONER EL MES AQUI ####
###########################
MES_ACTUAL = "Enero"
###########################

# Configurar el idioma local a español
locale.setlocale(locale.LC_TIME, 'es_ES')

# Crear un nuevo documento
document = Document()

###########################################################################################################################################
########################################################### ENCABEZADO ####################################################################
###########################################################################################################################################

# TITULO
titulo = document.add_paragraph('INFORME DE SERVICIOS')
titulo.runs[0].bold = True
titulo.runs[0].font.name = 'Arial'
titulo.runs[0].font.size = Pt(13)
titulo.alignment = 1
titulo.paragraph_format.space_after = Pt(0)

# SUBTITULO
now = datetime.now()
mes = now.month
year = now.year
# Obtener el nombre del mes
mes_nombre = calendar.month_name[mes]
# Obtener el primer día del mes actual
primer_dia = datetime(year, mes, 1)
# Obtener el último día del mes actual
ultimo_dia = datetime(year, mes, calendar.monthrange(year, mes)[1])
# Obtener el último día hábil del mes actual
ultimo_dia_habil = primer_dia.replace(day=calendar.monthrange(year, mes)[1])
while ultimo_dia_habil.weekday() >= 5:  # Si es sábado o domingo, retrocede un día
    ultimo_dia_habil -= timedelta(days=1)

titulo2 = document.add_paragraph('CORRESPONDIENTE AL MES DE ' + str(mes_nombre).upper() + ' DE ' + str(year).upper())
titulo2.runs[0].font.name = 'Arial'
titulo2.runs[0].font.size = Pt(13)
titulo2.alignment = 1
titulo2.paragraph_format.space_after = Pt(0)

# ESPACIO
espacio = document.add_paragraph('').paragraph_format.space_after = Pt(0)

# Agregar una tabla de 9x4
table = document.add_table(rows=9, cols=4, style='Table Grid')

# Configurar el ancho de las columnas
anchos_columnas = [30, 185, 108, 107]
for i, ancho in enumerate(anchos_columnas):
    column = table.columns[i]
    for cell in column.cells:
        cell.width = Pt(ancho)

# Rellenar la tabla con contenido
contenido = [
    ['1.', 'PARA:', 'Gerencia de informática', ''],
    ['2.', 'DE:', 'Cristopher Abrahan Paiz López', ''],
    ['3.', 'NÚMERO DE CONTRATO:', 'SAT-GRRHH-040-2025', ''],
    ['4.', 'NÚMERO DE IDENTIFICACIÓN TRIBUTARIA NIT:', '11115133-3', ''],
    ['5.', 'RENGLÓN PRESUPUESTARIO:', '029 “Otras remuneraciones de personal temporal”', ''],
    ['6.', 'TIPO DE SERVICIOS PRESTADOS:', 'Servicios Técnicos', ''],
    ['7.', 'VIGENCIA DE CONTRATO:', 'DEL: 02/01/2025', 'AL: 31/12/2025'],
    ['8.', 'PERIODO DE ESTE INFORME:', 'DEL: ' + primer_dia.strftime("%d/%m/%Y"), 'AL: ' + ultimo_dia.strftime("%d/%m/%Y")],
    ['9.', 'LUGAR Y FECHA:', ultimo_dia_habil.strftime("Guatemala, %d de %B de %Y"), ''],
]

# Configurar el formato de fuente y tamaño para toda la tabla
for i, row_content in enumerate(contenido):
    for j, cell_content in enumerate(row_content):
        cell = table.cell(i, j)
        cell.paragraphs[0].paragraph_format.space_after = Pt(21)  # Establecer espacio después a cero
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if i in [0, 1, 2, 3, 4, 5, 8]:  # Filas a fusionar
            cell.paragraphs[0].paragraph_format.space_after = Pt(21)  # Restablecer espacio después

        cell.paragraphs[0].paragraph_format.space_before = Pt(21)
        cell.paragraphs[0].alignment = 1 if j == 0 else 0  # Alineación centrada para la columna 0, izquierda para las otras
        cell.paragraphs[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = cell.paragraphs[0].add_run(cell_content)
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        # Unir celdas según las especificaciones
        if j == 2 and i in [0, 1, 2, 3, 4, 5, 8]:
            cell.merge(table.cell(i, j+1))

###########################################################################################################################################
########################################################### ACTIVIDADES ###################################################################
###########################################################################################################################################

# ESPACIO
espacio = document.add_paragraph('').paragraph_format.space_after = Pt(0)
espacio = document.add_paragraph('').paragraph_format.space_after = Pt(0)

def corregir_ruta(ruta):
    return ruta.replace("\\", "/")

# Agregar una tabla de 2x1 con un encabezado que diga "Actividades"
table = document.add_table(rows=1, cols=1, style='Table Grid')
table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(5)
table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(5)
table.cell(0, 0).paragraphs[0].alignment = 1
table.cell(0, 0).paragraphs[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
run = table.cell(0, 0).paragraphs[0].add_run('ACTIVIDADES REALIZADAS')
run.font.name = 'Arial'
run.font.size = Pt(12)
table.cell(0, 0).paragraphs[0].runs[0].bold = True

# Carpeta de archivos HTML
ruta_carpeta_html = f"C:/Users/acapaizlo/Desktop/{MES_ACTUAL}"

# Carpeta de imágenes
ruta_carpeta_imagenes = f"C:/Users/acapaizlo/Pictures/Screenshots/2025/{MES_ACTUAL}"

# Iterar sobre los archivos HTML
for archivo_html in os.listdir(ruta_carpeta_html):
    ruta_archivo_html = os.path.join(ruta_carpeta_html, archivo_html)

    # Leer el contenido del archivo HTML
    with open(ruta_archivo_html, 'r', encoding='utf-8') as archivo:
        contenido_html = archivo.read()

    # Analizar el HTML con BeautifulSoup
    soup = BeautifulSoup(contenido_html, 'html.parser')

    # Obtener el título
    titulo_elemento = soup.find('h2', class_='issue-detail-header')
    if titulo_elemento:
        titulo_texto = titulo_elemento.text.strip()

        nueva_fila = table.add_row().cells
        nueva_fila[0].text = f'{titulo_texto}\n'
        titulo_run = nueva_fila[0].paragraphs[0].runs[0]
        titulo_run.font.name = 'Arial'
        titulo_run.font.size = Pt(12)
        titulo_run.font.underline = False  # Quitar subrayado
        titulo_run.font.bold = True  # Texto en negrita
        titulo_run.font.highlight_color = 7  # Resaltado amarillo
        nueva_fila[0].paragraphs[0].runs[0].text = nueva_fila[0].paragraphs[0].runs[0].text.replace("\n", "")
        nueva_fila[0].paragraphs[0].paragraph_format.space_after = Pt(7)
        nueva_fila[0].paragraphs[0].paragraph_format.space_before = Pt(7)
        nueva_fila[0].paragraphs[0].runs[0].text = nueva_fila[0].paragraphs[0].runs[0].text.upper()


    # Iterar sobre los divs de subtareas
    divs_subtareas = soup.find_all(lambda tag: tag.name == 'div' and 'journal' in tag.get('class', []))

    # Obtener el nombre corto del mes actual
    mes_actual = calendar.month_abbr[datetime.now().month]

    # Mapear nombres de los meses a números
    meses_mapa = {
        'Ene': 1,
        'Feb': 2,
        'Mar': 3,
        'Abr': 4,
        'May': 5,
        'Jun': 6,
        'Jul': 7,
        'Ago': 8,
        'Set': 9,
        'Oct': 10,
        'Nov': 11,
        'Dic': 12,
        'Sep': 9,
        'Oct': 10,
        'Dec': 12,
    }


    # Iterar sobre los divs de subtareas
    for div_subtarea in divs_subtareas:
        # Verificar si tiene la clase adicional 'has-details'
        tiene_has_details = 'has-details' in div_subtarea.get('class', [])
        # Obtener la fecha
        fecha_elemento = div_subtarea.find('h4')
        fecha_texto = fecha_elemento.text.split()[-5:-2]  # Tomar solo el día y el mes
        dia, mes_abr, anio = fecha_texto[0], fecha_texto[1], fecha_texto[2]

        # Convertir el nombre corto del mes a número
        mes_numero = meses_mapa.get(mes_abr)

        # Verificar si la fecha está en el mes actual
        if mes_numero and mes_numero == datetime.now().month:
            # Obtener el nombre de la imagen
            nombre_imagen_elemento = div_subtarea.find('a', title='Descargar')
            nombre_imagen = nombre_imagen_elemento.text if nombre_imagen_elemento else None

            # Obtener el texto de la subtarea
            texto_subtarea_elemento = div_subtarea.find('div', class_='wiki editable')
            texto_subtarea = texto_subtarea_elemento.text.strip() if texto_subtarea_elemento else None
            # Eliminar cadena no deseada
            texto_subtarea = texto_subtarea.replace("Otras opcionesCitarModificarCrear tema", "")

            # Agregar una nueva fila a la tabla
            nueva_fila = table.add_row().cells
            nueva_fila2 = table.add_row().cells

            # Agregar la fecha a la nueva fila
            nueva_fila[0].text = f'Detalles subtarea del: {dia} {mes_abr} {anio}\n'
            detalle_run = nueva_fila[0].paragraphs[0].runs[0]
            detalle_run.font.name = 'Arial'
            detalle_run.font.size = Pt(12)
            detalle_run.font.color.rgb = RGBColor(0, 127, 255)  # Celeste
            detalle_run.font.bold = True  # Texto en negrita
            nueva_fila[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar el texto
            nueva_fila[0].paragraphs[0].runs[0].text = nueva_fila[0].paragraphs[0].runs[0].text.replace("\n", "")
            nueva_fila[0].paragraphs[0].paragraph_format.space_after = Pt(7)
            nueva_fila[0].paragraphs[0].paragraph_format.space_before = Pt(7)


            # Agregar el texto de la subtarea a la nueva fila
            nueva_fila2[0].text += f'{texto_subtarea}\n'
            nueva_fila2[0].paragraphs[0].runs[0].text = nueva_fila2[0].paragraphs[0].runs[0].text.replace("\n", "")


            # Si hay nombre de imagen, buscarla y agregarla a la nueva fila
            if nombre_imagen:
                try:
                    ruta_imagen = corregir_ruta(os.path.join(ruta_carpeta_imagenes, nombre_imagen))

                    # Insertar la imagen como InlineShape en la nueva fila
                    paragraph = nueva_fila2[0].add_paragraph()
                    run = paragraph.add_run()
                    picture = run.add_picture(ruta_imagen, width=Inches(6.0))
                    picture.alignment = 1  # Centrar la imagen en la celda

                    # # Añadir una nueva línea después de la imagen
                    # nueva_fila2[0].add_paragraph('\n')
                    nueva_fila2[0].paragraphs[0].paragraph_format.space_after = Pt(7)
                    nueva_fila2[0].paragraphs[0].paragraph_format.space_before = Pt(7)
                except:
                    pass

# ESPACIOS PARA LA FIRMA
for i in range(0, 10):
    espacio = document.add_paragraph('').paragraph_format.space_after = Pt(0)

# Agregar una tabla de 2x2
table_firmas = document.add_table(rows=2, cols=2)

# Lista de textos para las firmas
textos_firmas = ['____________________________________', '______________________', 'Cristopher Abrahan Paiz López', 'Vo. Bo.']

# Configurar el formato de fuente y tamaño para toda la tabla de firmas
for i in range(2):
    for j in range(2):
        cell = table_firmas.cell(i, j)
        cell.paragraphs[0].paragraph_format.space_after = Pt(5)  # Establecer espacio después a 5 puntos
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(5)
        cell.paragraphs[0].alignment = 1  # Alineación centrada para todas las celdas
        cell.paragraphs[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Obtener el índice correspondiente en la lista de textos de firmas
        indice = i * 2 + j
        if indice < len(textos_firmas):
            run = cell.paragraphs[0].add_run(textos_firmas[indice])
            run.font.name = 'Arial'
            run.font.size = Pt(11)

# Guardar el documento en el escritorio
document.save(f'C:/Users/acapaizlo/Desktop/informe {MES_ACTUAL} 2025 - Cristopher Abrahan Paiz López.docx')
#abrir el documento
os.startfile(f'C:/Users/acapaizlo/Desktop/informe {MES_ACTUAL} 2025 - Cristopher Abrahan Paiz López.docx')